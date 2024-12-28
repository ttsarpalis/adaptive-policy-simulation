###################################################
# Streamlit-based UI for the Market Simulation
# Carbon Tax Inclusive + AI-based Analysis + Word Download
# Each user must provide their own OpenAI key (saved in session)
###################################################

import streamlit as st
import numpy as np
import pandas as pd
import statsmodels.api as sm
import matplotlib.pyplot as plt
import itertools
import logging
import io  # For in-memory file handling
from docx import Document  # For generating Word docs
from docx.shared import Inches

import openai  # For GPT-based analysis
from openai import OpenAI  # Hypothetical client wrapper

from tqdm import tqdm
from scipy.optimize import brentq, fminbound
from io import BytesIO

###################################################
# GLOBAL LOGGING SETUP
###################################################
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

###################################################
# GLOBAL CONSTANTS & PARAMETERS
###################################################
MARKET_STATES = ['Low', 'Medium', 'High']
ACTIONS = ['Stay Out', 'Enter', 'Exit']
NUM_ACTIONS = len(ACTIONS)

# Demand/Market
DEMAND_SCALE = 1000
DEMAND_ELASTICITY = 1.5

# Firm Parameters
FIRM_SIZES = ['Small', 'Medium', 'Large']
SIZE_EFFECTS = {
    'Small': {'production_capacity': 50, 'fixed_cost_multiplier': 0.8},
    'Medium': {'production_capacity': 100, 'fixed_cost_multiplier': 1.0},
    'Large': {'production_capacity': 150, 'fixed_cost_multiplier': 1.2}
}

BASE_FIXED_COST = 50
BASE_VARIABLE_COST = 0.5
BASE_MARGINAL_COST_SLOPE = 0.01
BETA_COST = 1.2  # exponent in the cost function

# Policy/Simulation Parameters
INITIAL_SUBSIDY = 80
INITIAL_CARBON_TAX = 20
WELFARE_LOW_THRESHOLD = 10000
WELFARE_HIGH_THRESHOLD = 20000

SHOCK_PROBABILITY = 0.1
SHOCK_MAGNITUDE = 5

TIME_STEPS = 20
NUM_FIRMS = 50

BASE_LEARNING_RATE = 0.1
BASE_DISCOUNT_FACTOR = 0.9
BASE_EPSILON = 0.1

###################################################
# NUMBA ACCELERATION (OPTIONAL)
###################################################
try:
    from numba import njit
    USE_NUMBA = True
except ImportError:
    USE_NUMBA = False

# PROFIT FUNCTIONS (carbon tax inclusive)
if USE_NUMBA:
    @njit
    def _firm_profit_numba(price, quantity, subsidy, carbon_tax,
                           variable_cost, marginal_cost_slope, fixed_cost):
        """
        Profit = (price + subsidy)*quantity 
                 - [fixed_cost + variable_cost*Q 
                    + marginal_cost_slope*Q^BETA_COST 
                    + carbon_tax*Q]
        """
        cost = (fixed_cost 
                + variable_cost * quantity 
                + marginal_cost_slope * quantity ** BETA_COST
                + carbon_tax * quantity)
        revenue = (price + subsidy) * quantity
        return revenue - cost

    @njit
    def _profit_function_numba(q, price, subsidy, carbon_tax,
                               variable_cost, marginal_cost_slope, fixed_cost):
        return -_firm_profit_numba(price, q, subsidy, carbon_tax,
                                   variable_cost, marginal_cost_slope, fixed_cost)
else:
    def _firm_profit_numba(price, quantity, subsidy, carbon_tax,
                           variable_cost, marginal_cost_slope, fixed_cost):
        cost = (fixed_cost 
                + variable_cost * quantity 
                + marginal_cost_slope * quantity ** BETA_COST
                + carbon_tax * quantity)
        revenue = (price + subsidy) * quantity
        return revenue - cost

    def _profit_function_numba(q, price, subsidy, carbon_tax,
                               variable_cost, marginal_cost_slope, fixed_cost):
        return -_firm_profit_numba(price, q, subsidy, carbon_tax,
                                   variable_cost, marginal_cost_slope, fixed_cost)


###################################################
# CLASS DEFINITIONS
###################################################
class Firm:
    def __init__(self, firm_id, size, variable_cost, marginal_cost_slope,
                 fixed_cost, production_capacity, learning_rate, epsilon):
        self.firm_id = firm_id
        self.size = size
        self.variable_cost = variable_cost
        self.marginal_cost_slope = marginal_cost_slope
        self.fixed_cost = fixed_cost
        self.production_capacity = production_capacity
        self.learning_rate = learning_rate
        self.epsilon = epsilon

        self.quantity = 0
        self.profit = 0
        self.q_table = np.zeros((len(MARKET_STATES), NUM_ACTIONS))

        # Track consecutive negative profit
        self.consecutive_negative_profits = 0

    def compute_profit(self, price, quantity, subsidy, carbon_tax):
        return _firm_profit_numba(price, quantity, subsidy, carbon_tax,
                                  self.variable_cost,
                                  self.marginal_cost_slope,
                                  self.fixed_cost)

    def optimal_quantity(self, price, subsidy, carbon_tax):
        def profit_func(q):
            return _profit_function_numba(q, price, subsidy, carbon_tax,
                                          self.variable_cost,
                                          self.marginal_cost_slope,
                                          self.fixed_cost)
        q_opt = fminbound(profit_func, 0, self.production_capacity, disp=0)
        return q_opt


class Market:
    def __init__(self, demand_scale=DEMAND_SCALE, demand_elasticity=DEMAND_ELASTICITY,
                 use_piecewise_demand=True):
        self.demand_scale = demand_scale
        self.demand_elasticity = demand_elasticity
        self.use_piecewise_demand = use_piecewise_demand

    def demand(self, price):
        if not self.use_piecewise_demand:
            return self.demand_scale * price ** (-self.demand_elasticity)

        if price < 10:
            Q = 2000 - 50 * price
        elif price < 30:
            Q = 1500 - 20 * (price - 10)
        else:
            Q = 1100 - 5 * (price - 30)
        return max(0, Q)

    def aggregate_supply(self, price, subsidy, carbon_tax, firms_in_market):
        total_supply = 0
        for firm in firms_in_market:
            q_i = firm.optimal_quantity(price, subsidy, carbon_tax)
            total_supply += q_i
        return total_supply

    def excess_demand(self, price, subsidy, carbon_tax, firms_in_market):
        Q_d = self.demand(price)
        Q_s = self.aggregate_supply(price, subsidy, carbon_tax, firms_in_market)
        return Q_d - Q_s

    def market_equilibrium(self, subsidy, carbon_tax, firms_in_market):
        if not firms_in_market:
            return None

        P_lower, P_upper = 1e-6, 1e4
        def eq_func(p):
            return self.excess_demand(p, subsidy, carbon_tax, firms_in_market)

        f_low = eq_func(P_lower)
        f_high = eq_func(P_upper)

        if f_low * f_high > 0:
            logging.error("Equilibrium function does not change sign => cannot solve.")
            return None

        try:
            p_eq = brentq(eq_func, P_lower, P_upper, xtol=1e-6, maxiter=500)
            return p_eq
        except ValueError as e:
            logging.error(f"Market eq not found: {e}")
            return None


class Policy:
    def __init__(self, subsidy=INITIAL_SUBSIDY, carbon_tax=INITIAL_CARBON_TAX,
                 shock_prob=SHOCK_PROBABILITY, shock_magnitude=SHOCK_MAGNITUDE,
                 government_budget=50000):
        self.subsidy = subsidy
        self.carbon_tax = carbon_tax
        self.shock_prob = shock_prob
        self.shock_magnitude = shock_magnitude
        self.government_budget = government_budget

        self.alpha_update = 0.5

        self.policy_history = {
            'Time_Step': [],
            'Subsidy': [],
            'Carbon_Tax': [],
            'Budget_Used': []
        }

    def record_policy(self, time_step, budget_used=0.0):
        self.policy_history['Time_Step'].append(time_step)
        self.policy_history['Subsidy'].append(self.subsidy)
        self.policy_history['Carbon_Tax'].append(self.carbon_tax)
        self.policy_history['Budget_Used'].append(budget_used)

    def apply_shocks(self, time_step):
        shock_occurred = False
        if np.random.rand() < self.shock_prob:
            shock_occurred = True
            if np.random.rand() < 0.5:
                shock_val = self.shock_magnitude * np.random.choice([-1, 1])
                new_subsidy = max(self.subsidy + shock_val, 0)
                self.subsidy = self.subsidy + self.alpha_update*(new_subsidy - self.subsidy)
                logging.info(f"External Shock at T={time_step}: partial update => subsidy={self.subsidy:.2f}")
            else:
                shock_val = self.shock_magnitude * np.random.choice([-1, 1])
                new_tax = max(self.carbon_tax + shock_val, 0)
                self.carbon_tax = self.carbon_tax + self.alpha_update*(new_tax - self.carbon_tax)
                logging.info(f"External Shock at T={time_step}: partial update => carbon_tax={self.carbon_tax:.2f}")
        return shock_occurred

    def standard_policy_update(self, time_step, welfare):
        subsidy_change = 0
        carbon_tax_change = 0

        if welfare < WELFARE_LOW_THRESHOLD:
            subsidy_change += self.shock_magnitude
            carbon_tax_change -= self.shock_magnitude
        elif welfare > WELFARE_HIGH_THRESHOLD:
            subsidy_change -= self.shock_magnitude
            carbon_tax_change += self.shock_magnitude

        new_subsidy = max(self.subsidy + subsidy_change, 0)
        self.subsidy = self.subsidy + self.alpha_update*(new_subsidy - self.subsidy)

        new_tax = max(self.carbon_tax + carbon_tax_change, 0)
        self.carbon_tax = self.carbon_tax + self.alpha_update*(new_tax - self.carbon_tax)

    def advanced_policy_feedback(self, time_step, welfare, total_production):
        spending = self.subsidy * total_production
        budget_excess = spending - self.government_budget

        if budget_excess > 0:
            cut = 0.05 * budget_excess
            new_subsidy = max(self.subsidy - cut, 0)
            self.subsidy = self.subsidy + self.alpha_update*(new_subsidy - self.subsidy)

        if welfare < (0.8 * WELFARE_LOW_THRESHOLD):
            boost = 0.1 * (WELFARE_LOW_THRESHOLD - welfare) / WELFARE_LOW_THRESHOLD
            new_subsidy = self.subsidy + boost * self.shock_magnitude
            self.subsidy = self.subsidy + self.alpha_update*(new_subsidy - self.subsidy)

        if welfare > (1.2 * WELFARE_HIGH_THRESHOLD):
            reduce_amt = 0.1 * (welfare - WELFARE_HIGH_THRESHOLD) / WELFARE_HIGH_THRESHOLD
            new_subsidy = max(self.subsidy - reduce_amt * self.shock_magnitude, 0)
            self.subsidy = self.subsidy + self.alpha_update*(new_subsidy - self.subsidy)

    def update_policies(self, time_step, welfare, total_production, use_advanced=True):
        self.standard_policy_update(time_step, welfare)
        if use_advanced:
            self.advanced_policy_feedback(time_step, welfare, total_production)


class Simulation:
    def __init__(self, firms, market, policy, num_time_steps=TIME_STEPS,
                 base_discount=BASE_DISCOUNT_FACTOR):
        self.firms = firms
        self.market = market
        self.policy = policy
        self.num_time_steps = num_time_steps
        self.discount_factor = base_discount

        self.firms_in_market = [f for f in self.firms if np.random.rand() < 0.5]
        self.welfare_over_time = []
        self.simulation_summary = {}

    def assign_market_state(self, price):
        if price > 50:
            return 'High'
        elif price > 20:
            return 'Medium'
        else:
            return 'Low'

    def choose_action(self, firm, state_idx):
        if np.random.rand() < firm.epsilon:
            return np.random.choice(NUM_ACTIONS)
        else:
            q_vals = firm.q_table[state_idx]
            max_q = np.max(q_vals)
            candidates = np.where(q_vals == max_q)[0]
            return np.random.choice(candidates)

    def update_q_table(self, firm, state_idx, action, reward, next_state_idx):
        current_q = firm.q_table[state_idx, action]
        max_future_q = np.max(firm.q_table[next_state_idx])
        firm.q_table[state_idx, action] += firm.learning_rate * \
            (reward + self.discount_factor * max_future_q - current_q)

    def compute_welfare(self, price):
        """
        W = Consumer Surplus + Producer Surplus - Subsidy Cost
        """
        if DEMAND_ELASTICITY == 1:
            CS = DEMAND_SCALE * np.log(price)
        else:
            CS = (DEMAND_SCALE / (DEMAND_ELASTICITY - 1)) * price ** (1 - DEMAND_ELASTICITY)

        PS = 0
        total_Q = 0
        for firm in self.firms_in_market:
            Q_i = firm.quantity
            PS += firm.profit  # profit includes carbon tax deduction
            total_Q += Q_i

        SC = self.policy.subsidy * total_Q
        return CS + PS - SC, total_Q

    def run(self):
        prev_price = 50
        for t in range(self.num_time_steps):
            self.policy.apply_shocks(time_step=t+1)

            if t > 0 and self.welfare_over_time:
                prev_welfare = self.welfare_over_time[-1]
                _, last_Q = self.compute_welfare(prev_price)
                self.policy.update_policies(t+1, prev_welfare, last_Q, use_advanced=True)

            # Firms choose actions
            current_state = self.assign_market_state(prev_price)
            state_idx = MARKET_STATES.index(current_state)
            action_records = {}
            for firm in self.firms:
                action = self.choose_action(firm, state_idx)
                action_records[firm.firm_id] = action
                action_name = ACTIONS[action]
                if action_name == 'Enter' and firm not in self.firms_in_market:
                    self.firms_in_market.append(firm)
                    firm.quantity = 0
                elif action_name == 'Exit' and firm in self.firms_in_market:
                    self.firms_in_market.remove(firm)
                    firm.quantity = 0

            # Solve equilibrium
            price_eq = self.market.market_equilibrium(
                self.policy.subsidy, 
                self.policy.carbon_tax,
                self.firms_in_market
            )
            if price_eq is None:
                if not self.firms_in_market:
                    price_eq = 100.0
                else:
                    logging.warning("No equilibrium found. Using fallback=200.0")
                    price_eq = 200.0
            prev_price = price_eq

            # Update each firm's production, profit
            for firm in self.firms_in_market[:]:
                q_opt = firm.optimal_quantity(price_eq, 
                                              self.policy.subsidy,
                                              self.policy.carbon_tax)
                q_opt = min(q_opt, firm.production_capacity)
                firm.quantity = q_opt
                pft = firm.compute_profit(price_eq, q_opt,
                                          self.policy.subsidy,
                                          self.policy.carbon_tax)
                firm.profit = pft

                if pft < 0:
                    firm.consecutive_negative_profits += 1
                else:
                    firm.consecutive_negative_profits = 0
                if firm.consecutive_negative_profits >= 2:
                    self.firms_in_market.remove(firm)
                    firm.quantity = 0

            # Compute welfare
            welfare_t, total_Q_t = self.compute_welfare(price_eq)
            self.welfare_over_time.append(welfare_t)

            # Record policy
            budget_used = self.policy.subsidy * total_Q_t
            self.policy.record_policy(time_step=t+1, budget_used=budget_used)

            # Q-learning updates
            next_state = self.assign_market_state(price_eq)
            next_state_idx = MARKET_STATES.index(next_state)
            for firm in self.firms:
                action = action_records[firm.firm_id]
                action_name = ACTIONS[action]
                if action_name == 'Enter':
                    if firm in self.firms_in_market:
                        reward = firm.profit if firm.profit >= 0 else -10
                    else:
                        reward = -10
                elif action_name == 'Exit':
                    if firm not in self.firms_in_market:
                        reward = -5
                    else:
                        reward = 0
                else:
                    if firm in self.firms_in_market:
                        reward = firm.profit if firm.profit >= 0 else -10
                    else:
                        reward = 0

                self.update_q_table(firm, state_idx, action, reward, next_state_idx)

        self.simulation_summary = {
            'Final_Social_Welfare': self.welfare_over_time[-1] if self.welfare_over_time else None,
            'Average_Social_Welfare': np.mean(self.welfare_over_time) if self.welfare_over_time else None,
            'Max_Social_Welfare': np.max(self.welfare_over_time) if self.welfare_over_time else None,
            'Min_Social_Welfare': np.min(self.welfare_over_time) if self.welfare_over_time else None,
            'Policy_History': self.policy.policy_history,
            'Welfare_Over_Time': self.welfare_over_time
        }
        return self.simulation_summary


###################################################
# UTILITY FUNCTIONS
###################################################
def generate_firms(num_firms=NUM_FIRMS):
    size_distribution = np.random.choice(FIRM_SIZES, size=num_firms, p=[0.2, 0.5, 0.3])
    firms = []
    for i in range(num_firms):
        size = size_distribution[i]
        capacity = SIZE_EFFECTS[size]['production_capacity']
        fmult = SIZE_EFFECTS[size]['fixed_cost_multiplier']

        fixed_cost = BASE_FIXED_COST * fmult
        variable_cost = BASE_VARIABLE_COST * np.random.uniform(0.9, 1.1)
        mcslope = BASE_MARGINAL_COST_SLOPE * np.random.uniform(0.9, 1.1)
        lr = np.random.uniform(0.05, 0.2)
        eps = np.random.uniform(0.05, 0.2)

        firm_obj = Firm(firm_id=i,
                        size=size,
                        variable_cost=variable_cost,
                        marginal_cost_slope=mcslope,
                        fixed_cost=fixed_cost,
                        production_capacity=capacity,
                        learning_rate=lr,
                        epsilon=eps)
        firms.append(firm_obj)
    return firms

###################################################
# PLOTTING HELPERS
###################################################
def plot_welfare(sim_summary):
    fig, ax = plt.subplots(figsize=(8,5))
    welf_list = sim_summary['Welfare_Over_Time']
    ax.plot(range(1, len(welf_list)+1), welf_list, marker='o')
    ax.set_title('Welfare Over Time')
    ax.set_xlabel('Time Step')
    ax.set_ylabel('Social Welfare')
    ax.grid(True)
    return fig

def plot_policy(sim_summary):
    pol_hist = sim_summary['Policy_History']
    steps = pol_hist['Time_Step']

    fig, ax = plt.subplots(figsize=(8,5))

    line1 = ax.plot(steps, pol_hist['Subsidy'], label='Subsidy', color='blue')
    line2 = ax.plot(steps, pol_hist['Carbon_Tax'], label='Carbon Tax', color='orange')
    ax.set_xlabel('Time Step')
    ax.set_ylabel('Subsidy / Carbon Tax')

    ax2 = ax.twinx()
    line3 = ax2.plot(steps, pol_hist['Budget_Used'], label='Budget Used', color='green')
    ax2.set_ylabel('Budget Used')

    lines = line1 + line2 + line3
    labels = [l.get_label() for l in lines]
    ax2.legend(lines, labels, loc='best')

    ax.set_title('Policy Trends (Separate Scales)')
    ax.grid(True)
    fig.tight_layout()
    return fig

###################################################
# AI ANALYSIS FUNCTIONS
###################################################
def run_gpt_analysis(sim_results, user_api_key):
    """
    Requires each user to provide their own key. We pass user_api_key in.
    """
    # Create a client with the user's key, not a hardcoded one.
    client = OpenAI(api_key=user_api_key)

    prompt_text = (
        "You are a helpful AI analyzing a market simulation. "
        "Here are the final results:\n"
        f"Final Welfare = {sim_results['Final_Social_Welfare']}\n"
        f"Average Welfare = {sim_results['Average_Social_Welfare']}\n"
        f"Max Welfare = {sim_results['Max_Social_Welfare']}\n"
        f"Min Welfare = {sim_results['Min_Social_Welfare']}\n\n"
        "Please provide a concise analysis of these results, "
        "highlighting any interesting behavior or potential expansions."
    )

    response = client.chat.completions.create(
        messages=[
            {
                "role": "user",
                "content": prompt_text,
            }
        ],
        model="gpt-4o-mini",
    )

    # For debugging, let's see if we get a request id:
    if hasattr(response, "_request_id"):
        print(f"Request ID: {response._request_id}")

    if hasattr(response, 'choices') and len(response.choices) > 0:
        analysis_text = response.choices[0].message.content
    else:
        analysis_text = "No analysis returned by GPT-4o-mini."

    return analysis_text

def create_word_doc_with_analysis(analysis_text, fig_welfare, fig_policy):
    """
    Creates an in-memory .docx file with the AI analysis text and the two charts.
    """
    doc = Document()
    doc.add_heading('AI Analysis of Simulation', level=1)

    # Add analysis text
    doc.add_paragraph(analysis_text)

    # Convert matplotlib figs to images
    buf_welfare = io.BytesIO()
    fig_welfare.savefig(buf_welfare, format='png')
    buf_welfare.seek(0)

    buf_policy = io.BytesIO()
    fig_policy.savefig(buf_policy, format='png')
    buf_policy.seek(0)

    doc.add_heading('Welfare Over Time Chart', level=2)
    doc.add_picture(buf_welfare, width=Inches(5.5))

    doc.add_heading('Policy Trends Chart', level=2)
    doc.add_picture(buf_policy, width=Inches(5.5))

    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

###################################################
# STREAMLIT APP
###################################################
def main():
    st.title("Adaptive Policy Simulation: Carbon Pricing and Firm Behavior in a Partial-Equilibrium Framework")

    st.sidebar.header("Simulation Controls")
    # First, ask for user to supply their own API key.
    user_api_key = st.sidebar.text_input("Enter your OpenAI API key:", value="", type="password")

    num_firms = st.sidebar.slider("Number of Firms", 10, 200, 50, step=10)
    time_steps = st.sidebar.slider("Time Steps", 5, 50, 20, step=5)
    subsidy_init = st.sidebar.slider("Initial Subsidy", 0, 200, 80, step=10)
    carbon_tax_init = st.sidebar.slider("Initial Carbon Tax", 0, 100, 20, step=5)
    gov_budget = st.sidebar.slider("Gov Budget Constraint", 10000, 200000, 50000, step=5000)
    shock_prob = st.sidebar.slider("Shock Probability", 0.0, 1.0, 0.1, step=0.1)
    shock_magnitude = st.sidebar.slider("Shock Magnitude", 0, 50, 5, step=5)

    # Session state variables
    if "sim_results" not in st.session_state:
        st.session_state.sim_results = None
    if "fig_welfare" not in st.session_state:
        st.session_state.fig_welfare = None
    if "fig_policy" not in st.session_state:
        st.session_state.fig_policy = None
    if "analysis_text" not in st.session_state:
        st.session_state.analysis_text = None
    if "doc_file" not in st.session_state:
        st.session_state.doc_file = None

    if st.button("Run Simulation"):
        firms = generate_firms(num_firms=num_firms)
        market = Market(use_piecewise_demand=True)
        policy = Policy(
            subsidy=subsidy_init,
            carbon_tax=carbon_tax_init,
            shock_prob=shock_prob,
            shock_magnitude=shock_magnitude,
            government_budget=gov_budget
        )

        sim = Simulation(firms=firms, market=market, policy=policy, num_time_steps=time_steps)
        sim_results = sim.run()
        st.session_state.sim_results = sim_results

        st.subheader("Simulation Results")
        if sim_results['Welfare_Over_Time']:
            st.write(f"**Final Welfare:** {sim_results['Final_Social_Welfare']:.2f}")
            st.write(f"**Average Welfare:** {sim_results['Average_Social_Welfare']:.2f}")
            st.write(f"**Max Welfare:** {sim_results['Max_Social_Welfare']:.2f}")
            st.write(f"**Min Welfare:** {sim_results['Min_Social_Welfare']:.2f}")

            fig_wel = plot_welfare(sim_results)
            st.pyplot(fig_wel)

            fig_pol = plot_policy(sim_results)
            st.pyplot(fig_pol)

            st.session_state.fig_welfare = fig_wel
            st.session_state.fig_policy = fig_pol
        else:
            st.write("No welfare data generated or 0 time steps.")

    # Only allow "Analyze" if user has a key and sim_results is ready.
    if st.session_state.sim_results is not None:
        if user_api_key.strip() == "":
            st.warning("Please enter your OpenAI API key in the sidebar to use AI analysis.")
        else:
            if st.button("Analyze"):
                # Now we call run_gpt_analysis with the user's key.
                analysis = run_gpt_analysis(st.session_state.sim_results, user_api_key)
                st.session_state.analysis_text = analysis
                st.write("**AI Analysis**:")
                st.markdown(analysis)

                if st.session_state.fig_welfare is not None and st.session_state.fig_policy is not None:
                    doc_file = create_word_doc_with_analysis(
                        analysis,
                        st.session_state.fig_welfare,
                        st.session_state.fig_policy
                    )
                    st.session_state.doc_file = doc_file
                else:
                    st.write("No figures found to embed in the Word doc.")

    if st.session_state.doc_file is not None:
        st.download_button(
            label="Download Analysis (Word doc)",
            data=st.session_state.doc_file,
            file_name="simulation_analysis.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

if __name__ == "__main__":
    main()
